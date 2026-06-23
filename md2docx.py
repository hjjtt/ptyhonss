# -*- coding: utf-8 -*-
"""通用 Markdown -> Word 转换器
用法: python md2docx.py <源.md> <目标.docx>
保留: 标题层级 / 表格 / 代码块 / 列表 / 引用 / 段落 / 行内代码 / 粗体斜体
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------- 样式常量 ----------------
CN_FONT = "Microsoft YaHei"
EN_FONT = "Calibri"
CODE_FONT = "Consolas"

COLOR_H1 = RGBColor(0x1F, 0x4E, 0x79)
COLOR_H2 = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H3 = RGBColor(0x44, 0x72, 0xC4)
COLOR_CODE = RGBColor(0x35, 0x35, 0x35)
COLOR_QUOTE = RGBColor(0x60, 0x60, 0x60)
COLOR_INLINE_CODE = RGBColor(0xC7, 0x25, 0x4E)


def set_cn_font(run, font_name=CN_FONT):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_runs_with_inline(paragraph, text, base_size=10.5):
    pattern = re.compile(r"(`[^`]+`)|(\*\*[^*]+\*\*)|(\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_cn_font(run)
            run.font.size = Pt(base_size)
        seg = m.group()
        if seg.startswith("`"):
            run = paragraph.add_run(seg[1:-1])
            set_cn_font(run, CODE_FONT)
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = COLOR_INLINE_CODE
        elif seg.startswith("**"):
            run = paragraph.add_run(seg[2:-2])
            set_cn_font(run)
            run.font.size = Pt(base_size)
            run.bold = True
        else:
            run = paragraph.add_run(seg[1:-1])
            set_cn_font(run)
            run.font.size = Pt(base_size)
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_cn_font(run)
        run.font.size = Pt(base_size)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    tc_pr.append(shd)


def add_heading(doc, text, level):
    sizes = {1: 20, 2: 16, 3: 13.5, 4: 12}
    colors = {1: COLOR_H1, 2: COLOR_H2, 3: COLOR_H3, 4: COLOR_H3}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_cn_font(run)
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = colors.get(level, COLOR_H3)
    run.bold = True
    if level == 1:
        p.paragraph_format.space_before = Pt(20)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "6",
            qn("w:space"): "4", qn("w:color"): "2E74B5",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F5F5F5",
    })
    pPr.append(shd)
    run = p.add_run("\n".join(code_lines))
    set_cn_font(run, CODE_FONT)
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_CODE


def add_table(doc, header, rows):
    n_col = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_col)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs_with_inline(p, h, base_size=10)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "2E74B5")
    for r_idx, row in enumerate(rows, 1):
        for c_idx in range(n_col):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_with_inline(p, row[c_idx] if c_idx < len(row) else "", base_size=10)
            if r_idx % 2 == 0:
                shade_cell(cell, "F2F7FC")


def convert(src_path, dst_path):
    md = Path(src_path).read_text(encoding="utf-8")
    lines = md.split("\n")
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # 代码块
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and stripped.endswith("|"):
            if i + 1 < n and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
                def parse_row(row):
                    return [c.strip() for c in row.strip().strip("|").split("|")]
                header = parse_row(stripped)
                i += 2
                rows = []
                while i < n and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    rows.append(parse_row(lines[i].strip()))
                    i += 1
                add_table(doc, header, rows)
                continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # 跳过目录块
            if text.strip() == "目录":
                while i < n and stripped != "---":
                    i += 1
                    if i < n:
                        stripped = lines[i].rstrip()
                i += 1
                continue
            add_heading(doc, text, level)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            if not text:
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            left = pBdr.makeelement(qn("w:left"), {
                qn("w:val"): "single", qn("w:sz"): "18",
                qn("w:space"): "8", qn("w:color"): "2E74B5",
            })
            pBdr.append(left)
            pPr.append(pBdr)
            shd = pPr.makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F2F7FC",
            })
            pPr.append(shd)
            add_runs_with_inline(p, text, base_size=10)
            for run in p.runs:
                run.font.color.rgb = COLOR_QUOTE
            i += 1
            continue

        # 无序列表
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent_level = len(m.group(1)) // 2
            text = m.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.6)
            p.paragraph_format.space_after = Pt(2)
            add_runs_with_inline(p, text, base_size=10.5)
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            indent_level = len(m.group(1)) // 2
            text = m.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.6)
            p.paragraph_format.space_after = Pt(2)
            add_runs_with_inline(p, text, base_size=10.5)
            i += 1
            continue

        # 水平分隔线
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:space"): "1", qn("w:color"): "BFBFBF",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.4
        add_runs_with_inline(p, stripped, base_size=10.5)
        i += 1

    doc.save(dst_path)
    return Path(dst_path).stat().st_size / 1024


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python md2docx.py <源.md> <目标.docx>")
        sys.exit(1)
    size = convert(sys.argv[1], sys.argv[2])
    print(f"已生成: {Path(sys.argv[2]).resolve()} ({size:.1f} KB)")
